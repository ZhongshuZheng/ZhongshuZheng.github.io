import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches

from tools.docx_to_markdown import convert


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class DocxToMarkdownTests(unittest.TestCase):
    def test_converts_document_structure_and_image(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            source = root / "示例笔记.docx"
            output = root / "note" / "index.md"
            image = root / "fixture.png"
            image.write_bytes(ONE_PIXEL_PNG)

            document = Document()
            document.add_paragraph("示例笔记")
            document.add_heading("章节一", level=1)
            document.add_paragraph("第一项", style="List Bullet")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "名称"
            table.cell(0, 1).text = "年代"
            table.cell(1, 0).text = "示例"
            table.cell(1, 1).text = "1200"
            document.add_picture(str(image), width=Inches(1))
            document.save(source)

            report = convert(source, output)
            markdown = output.read_text(encoding="utf-8")

            self.assertIn("# 示例笔记", markdown)
            self.assertIn("## 章节一", markdown)
            self.assertIn("- 第一项", markdown)
            self.assertIn("| 名称 | 年代 |", markdown)
            self.assertIn("images/image-01.png", markdown)
            self.assertTrue((output.parent / "images" / "image-01.png").exists())
            self.assertEqual(report["images"], 1)
            self.assertEqual(report["tables"], 1)


if __name__ == "__main__":
    unittest.main()
